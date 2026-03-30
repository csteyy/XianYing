#!/usr/bin/env python3
"""Generate thesis DOCX from Markdown source per SZU formatting spec.

Generates: Cover -> Declaration -> TOC -> Title/Abstract/Body/Refs/Thanks/EN Abstract -> Appendix.
Cover, declaration, and appendix are copied from official templates with fields filled in.
"""

import os, re, sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION_START

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
DIR = Path(__file__).resolve().parent
MD_FILE = DIR / "论文初稿_曹思越_AI编程范式驱动的数字传播产品开发实践.md"
OUT_FILE = DIR / "论文初稿_曹思越_AI编程范式驱动的数字传播产品开发实践.docx"

BLACK = RGBColor(0, 0, 0)

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_KAI = "楷体"
FONT_TNR = "Times New Roman"

PT_1 = Pt(26)       # 一号
PT_XIAO1 = Pt(24)   # 小一号
PT_XIAO2 = Pt(18)   # 小二号
PT_3 = Pt(16)       # 三号
PT_XIAO3 = Pt(15)   # 小三号
PT_4 = Pt(14)       # 四号
PT_XIAO4 = Pt(12)   # 小四号
PT_5 = Pt(10.5)     # 五号
PT_XIAO5 = Pt(9)    # 小五号

FONT_HWZS = "华文中宋"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_run_font(run, name, size, bold=False, east_asia=None, color=BLACK):
    """Set font properties on a run, including cs_size for complex scripts."""
    rpr = run._r.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    ea = east_asia or name
    fonts.set(qn("w:eastAsia"), ea)
    if name == FONT_TNR:
        fonts.set(qn("w:ascii"), FONT_TNR)
        fonts.set(qn("w:hAnsi"), FONT_TNR)
    else:
        fonts.set(qn("w:ascii"), ea)
        fonts.set(qn("w:hAnsi"), ea)
    fonts.set(qn("w:cs"), FONT_TNR)
    rpr.append(fonts)
    run.font.size = size
    run.font.cs_size = size
    run.font.bold = bold
    run.font.color.rgb = color


def set_paragraph_spacing(para, before_lines=50, after_lines=50,
                          line_spacing=1.0):
    """Set paragraph spacing using line-based units (beforeLines/afterLines).

    before_lines / after_lines: in 1/100 of a line (50 = 0.5 line).
    line_spacing: multiplier (1.0 = single, 1.5 = 1.5x).
    """
    para.paragraph_format.line_spacing = line_spacing
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:beforeLines"), str(before_lines))
    spacing.set(qn("w:afterLines"), str(after_lines))
    line_val = str(int(240 * line_spacing))
    spacing.set(qn("w:line"), line_val)
    spacing.set(qn("w:lineRule"), "auto")


def set_first_line_indent(para, chars=200):
    """Set first-line indent using character count (200 = 2 chars)."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars))


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def set_page_layout(doc):
    """A4 page, margins per SZU spec (2.5cm top/bottom, 2.5cm left/right)."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ---------------------------------------------------------------------------
# Citation handling: [n] -> superscript
# ---------------------------------------------------------------------------
CITE_RE = re.compile(r"(\[\d+\])")


def add_runs_with_citations(para, text, font_name, font_size,
                            bold=False, east_asia=None):
    """Add text to paragraph, turning [n] citations into superscript runs."""
    parts = CITE_RE.split(text)
    for part in parts:
        if not part:
            continue
        run = para.add_run(part)
        if CITE_RE.match(part):
            set_run_font(run, font_name, PT_XIAO5, bold=False,
                         east_asia=east_asia)
            run.font.superscript = True
        else:
            set_run_font(run, font_name, font_size, bold=bold,
                         east_asia=east_asia)


# ---------------------------------------------------------------------------
# Table borders helper
# ---------------------------------------------------------------------------

def add_table_borders(table):
    """Add all borders and cell padding to a table via XML."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)
    # Cell padding: top/bottom 57 twips (~1mm) for vertical breathing room
    cell_mar = OxmlElement("w:tblCellMar")
    for side in ("top", "bottom"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "57")
        m.set(qn("w:type"), "dxa")
        cell_mar.append(m)
    tbl_pr.append(cell_mar)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


# ---------------------------------------------------------------------------
# MD Parser
# ---------------------------------------------------------------------------

def parse_md(path):
    """Parse the MD file into structured sections.

    Returns dict with keys:
        abstract, keywords, body_sections, references, thanks,
        en_abstract, en_keywords
    Each body_section is a list of items:
        ('h1', text), ('h2', text), ('h3', text),
        ('para', text), ('table', header_row, data_rows),
        ('note', text),
        ('mermaid', label_text, mermaid_code),
        ('fig4_html', html_block),
        ('img_label', text)
    """
    text = path.read_text(encoding="utf-8")
    lines = text.split("\n")

    result = {
        "title": "",
        "subtitle": "",
        "author": "",
        "affiliation": "",
        "advisor": "",
        "abstract": "",
        "keywords": "",
        "body_lines": [],
        "references": [],
        "thanks": "",
        "en_abstract": "",
        "en_keywords": "",
    }

    for j in range(min(15, len(lines))):
        ln = lines[j].strip()
        if ln.startswith("# ") and not ln.startswith("## "):
            result["title"] = ln[2:].strip()
        elif j > 0 and not result["subtitle"] and not ln.startswith("#") and ln and "为个案" in ln:
            result["subtitle"] = ln
        elif re.match(r'^[\u4e00-\u9fff]{2,4}$', ln):
            result["author"] = ln
        elif ln.startswith("\uff08") and "\u6df1\u5733\u5927\u5b66" in ln:
            result["affiliation"] = ln
        elif ln.startswith("\u6307\u5bfc\u6559\u5e08"):
            result["advisor"] = ln

    state = "skip"
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Detect abstract
        if stripped.startswith("【摘要】"):
            result["abstract"] = stripped[4:]
            state = "after_abstract"
            i += 1
            continue

        if state == "skip":
            i += 1
            continue

        if stripped.startswith("【关键词】"):
            result["keywords"] = stripped[5:]
            state = "body"
            i += 1
            continue

        if state == "after_abstract":
            i += 1
            continue

        # English abstract at end
        if stripped.startswith("【Abstract】"):
            result["en_abstract"] = stripped[len("【Abstract】"):]
            state = "en_abstract"
            i += 1
            continue

        if stripped.startswith("【Key words】") or stripped.startswith("【Key Words】"):
            kw_prefix = "【Key words】" if "【Key words】" in stripped else "【Key Words】"
            result["en_keywords"] = stripped[len(kw_prefix):]
            state = "done"
            i += 1
            continue

        if state == "en_abstract":
            i += 1
            continue

        if state == "done":
            i += 1
            continue

        # Body parsing
        if state == "body":
            # Separators
            if stripped == "---":
                i += 1
                continue

            # Heading ## X  ->  h1 (chapter)
            if stripped.startswith("## ") and not stripped.startswith("### "):
                title = stripped[3:].strip()
                if title == "参考文献":
                    state = "references"
                    i += 1
                    continue
                if title == "致谢":
                    state = "thanks"
                    i += 1
                    continue
                result["body_lines"].append(("h1", title))
                i += 1
                continue

            # Heading #### X  ->  h3 (must check before ###)
            if stripped.startswith("#### ") and not stripped.startswith("##### "):
                title = stripped[5:].strip()
                result["body_lines"].append(("h3", title))
                i += 1
                continue

            # Heading ### X  ->  h2
            if stripped.startswith("### ") and not stripped.startswith("#### "):
                title = stripped[4:].strip()
                result["body_lines"].append(("h2", title))
                i += 1
                continue

            # Mermaid code block
            if stripped.startswith("```mermaid"):
                mermaid_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    mermaid_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                result["body_lines"].append(("mermaid", "\n".join(mermaid_lines)))
                continue

            # HTML table block (figure 4)
            if stripped.startswith("<table"):
                html_lines = [line]
                i += 1
                while i < len(lines) and "</table>" not in lines[i]:
                    html_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    html_lines.append(lines[i])
                    i += 1
                result["body_lines"].append(("fig4_html", "\n".join(html_lines)))
                continue

            # Table (pipe-delimited)
            if stripped.startswith("|") and "|" in stripped[1:]:
                header = stripped
                i += 1
                if i < len(lines) and lines[i].strip().startswith("|"):
                    i += 1  # skip separator row
                rows = [header]
                while i < len(lines) and lines[i].strip().startswith("|"):
                    rows.append(lines[i].strip())
                    i += 1
                result["body_lines"].append(("table", rows))
                continue

            # Figure/table label lines like "图1 系统架构图" or "表1 ..."
            if re.match(r"^((?:图|表)\d+[a-z]?\s)", stripped):
                result["body_lines"].append(("img_label", stripped))
                i += 1
                continue

            # Parenthetical note
            if stripped.startswith("（注：") or stripped.startswith("(注："):
                result["body_lines"].append(("note", stripped))
                i += 1
                continue

            # Normal paragraph
            if stripped:
                result["body_lines"].append(("para", stripped))

            i += 1
            continue

        # References
        if state == "references":
            if stripped == "---":
                state = "body"
                i += 1
                continue
            if stripped.startswith("## 致谢"):
                state = "thanks"
                i += 1
                continue
            if stripped:
                result["references"].append(stripped)
            i += 1
            continue

        # Thanks
        if state == "thanks":
            if stripped == "---":
                state = "body"
                i += 1
                continue
            if stripped.startswith("【Abstract】"):
                result["en_abstract"] = stripped[len("【Abstract】"):]
                state = "en_abstract"
                i += 1
                continue
            if stripped:
                result["thanks"] += stripped + "\n"
            i += 1
            continue

        i += 1

    result["thanks"] = result["thanks"].strip()
    return result


# ---------------------------------------------------------------------------
# Image mapping: figure labels to PNG files
# ---------------------------------------------------------------------------
IMAGE_MAP = {
    1: "图1_系统架构.png",
    2: "图2_平台适配.png",
    3: "图3_数据流.png",
}


# ---------------------------------------------------------------------------
# Build functions
# ---------------------------------------------------------------------------

def build_title_page(doc, data):
    """Build title page: main title, subtitle, author, affiliation, advisor."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_lines=300, after_lines=0, line_spacing=1.5)
    run = p.add_run(data["title"])
    set_run_font(run, FONT_HWZS, PT_XIAO2, bold=True)

    if data.get("subtitle"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before_lines=0, after_lines=100, line_spacing=1.5)
        run2 = p2.add_run(data["subtitle"])
        set_run_font(run2, FONT_SONG, PT_XIAO4)

    if data.get("author"):
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p3, before_lines=50, after_lines=0)
        run3 = p3.add_run(data["author"])
        set_run_font(run3, FONT_SONG, PT_XIAO4)

    if data.get("affiliation"):
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p4, before_lines=0, after_lines=0)
        run4 = p4.add_run(data["affiliation"])
        set_run_font(run4, FONT_SONG, PT_XIAO4)

    if data.get("advisor"):
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p5, before_lines=0, after_lines=50)
        run5 = p5.add_run(data["advisor"])
        set_run_font(run5, FONT_SONG, PT_XIAO4)


def build_toc(doc):
    """Insert a dynamic TOC field."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_lines=0, after_lines=50)
    run = p.add_run("目录")
    set_run_font(run, FONT_SONG, PT_XIAO4, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p2, line_spacing=1.5)

    r1 = p2.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld_begin)

    r2 = p2.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)

    r3 = p2.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(fld_sep)

    toc_hint = '\uff08\u8bf7\u5728Word\u4e2d\u53f3\u952e\u6b64\u5904\uff0c\u9009\u62e9\u201c\u66f4\u65b0\u57df\u201d\u4ee5\u751f\u6210\u76ee\u5f55\uff09'
    r4 = p2.add_run(toc_hint)
    set_run_font(r4, FONT_SONG, PT_XIAO4)

    r5 = p2.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5._r.append(fld_end)


def build_abstract(doc, abstract_text, keywords_text):
    """Build Chinese abstract and keywords section."""
    # Abstract label + content
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p)

    label_run = p.add_run("【摘要】")
    set_run_font(label_run, FONT_KAI, PT_XIAO4, bold=True)
    add_runs_with_citations(p, abstract_text, FONT_KAI, PT_5)

    # Keywords label + content
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p2)

    kw_label = p2.add_run("【关键词】")
    set_run_font(kw_label, FONT_KAI, PT_XIAO4, bold=True)
    kw_run = p2.add_run(keywords_text)
    set_run_font(kw_run, FONT_KAI, PT_5)


def build_heading(doc, level, text):
    """Add a heading paragraph with correct SZU formatting.

    level 1 = ## (chapter), level 2 = ### (section), level 3 = #### (subsection)
    """
    p = doc.add_paragraph()

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before_lines=100, after_lines=50)
        run = p.add_run(text)
        set_run_font(run, FONT_HEI, PT_3, bold=True)
        p.style = doc.styles["Heading 1"]
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before_lines=50, after_lines=50)
        run = p.add_run(text)
        set_run_font(run, FONT_HEI, PT_XIAO3, bold=True)
        p.style = doc.styles["Heading 2"]
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before_lines=50, after_lines=50)
        run = p.add_run(text)
        set_run_font(run, FONT_HEI, PT_4, bold=True)
        p.style = doc.styles["Heading 3"]

    # Override style-inherited font/color
    for r in p.runs:
        r.font.color.rgb = BLACK


def build_body_paragraph(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p)
    set_first_line_indent(p)
    add_runs_with_citations(p, text, FONT_SONG, PT_5)


def build_note_paragraph(doc, text):
    """Add a parenthetical note paragraph (smaller font, no indent)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before_lines=25, after_lines=25)
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, PT_XIAO5)


def build_img_label(doc, text):
    """Add an image label paragraph (centered, bold, small)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_lines=25, after_lines=50)
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, PT_XIAO5, bold=True)


def build_image(doc, img_path, width=Inches(5)):
    """Insert an image centered in the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_lines=50, after_lines=25)
    run = p.add_run()
    run.add_picture(str(img_path), width=width)


def build_table(doc, rows):
    """Build a table from pipe-delimited rows."""
    def parse_row(r):
        cells = [c.strip() for c in r.split("|")]
        return [c for c in cells if c]

    header_cells = parse_row(rows[0])
    num_cols = len(header_cells)
    data_rows = [parse_row(r) for r in rows[1:]]

    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row with light blue-gray background
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = ""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E2F3")
        tc_pr.append(shd)
        v_align = OxmlElement("w:vAlign")
        v_align.set(qn("w:val"), "center")
        tc_pr.append(v_align)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_lines=0, after_lines=0, line_spacing=1.0)
        run = p.add_run(cell_text)
        set_run_font(run, FONT_SONG, PT_XIAO5, bold=True)

    # Data rows
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            tc_pr = cell._tc.get_or_add_tcPr()
            v_align = OxmlElement("w:vAlign")
            v_align.set(qn("w:val"), "center")
            tc_pr.append(v_align)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before_lines=0, after_lines=0, line_spacing=1.0)
            run = p.add_run(cell_text)
            set_run_font(run, FONT_SONG, PT_XIAO5)


def build_fig4(doc):
    """Build figure 4 as a borderless 2-column table."""
    img_a = DIR / "图4a_沙箱隔离视觉探索.png"
    img_b = DIR / "图4b_3D可视化效果.png"

    if not img_a.exists() or not img_b.exists():
        p = doc.add_paragraph("（图4图片文件缺失）")
        return

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: images (top)
    for j, img_path in enumerate([img_a, img_b]):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(2.8))

    # Row 1: captions (bottom)
    for j, caption in enumerate([
        "图4a 沙箱隔离环境中的视觉风格探索",
        "图4b 最终选定的三维互动可视化效果",
    ]):
        cell = table.rows[1].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        set_run_font(run, FONT_SONG, PT_XIAO5, bold=True)


def build_references(doc, ref_lines):
    """Build references section. Title: 五号楷体加粗顶格; Content: 小五号楷体顶格."""
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p_title)
    run_title = p_title.add_run("参考文献")
    set_run_font(run_title, FONT_KAI, PT_5, bold=True)

    for ref in ref_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p)
        run = p.add_run(ref)
        set_run_font(run, FONT_KAI, PT_XIAO5)


def build_thanks(doc, thanks_text):
    """Build acknowledgments section."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_lines=50, after_lines=50)
    run = p.add_run("致谢")
    set_run_font(run, FONT_HEI, PT_XIAO4, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p2)
    set_first_line_indent(p2)
    run2 = p2.add_run(thanks_text)
    set_run_font(run2, FONT_SONG, PT_5)


def build_en_abstract(doc, en_abstract, en_keywords):
    """Build English abstract section."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p)
    label = p.add_run("【Abstract】")
    set_run_font(label, FONT_TNR, PT_XIAO4, bold=True)
    body = p.add_run(en_abstract)
    set_run_font(body, FONT_TNR, PT_5)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p2)
    kw_label = p2.add_run("【Key words】")
    set_run_font(kw_label, FONT_TNR, PT_XIAO4, bold=True)
    kw_body = p2.add_run(en_keywords)
    set_run_font(kw_body, FONT_TNR, PT_5)


# ---------------------------------------------------------------------------
# Cover / Declaration / Appendix — generated from code
# ---------------------------------------------------------------------------

def _get_full_title(data):
    t = data["title"]
    if data.get("subtitle"):
        t += "\uff1a" + data["subtitle"]
    return t


def set_spacing_abs(para, before=0, after=0, line=None):
    """Set paragraph spacing using absolute twip values."""
    pPr = para._p.get_or_add_pPr()
    old = pPr.find(qn("w:spacing"))
    if old is not None:
        pPr.remove(old)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def set_first_line_indent_twips(para, twips):
    """Set first-line indent using absolute twip value."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLine"), str(twips))


def set_left_indent_twips(para, twips):
    """Set paragraph left indent in twips."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"), str(twips))


def _set_section_layout(section, top=2.54, bottom=2.54, left=3.17, right=3.17):
    """Configure A4 page layout on a specific section."""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def build_cover(doc, data):
    """Generate cover page (pixel-level replica of official template).

    Template margins: T=2.54cm B=2.54cm L=3.17cm R=3.17cm
    Font sizes extracted from template XML:
      title line = 26pt bold, subtitle = 24pt bold, fields = 16pt bold
    """
    full_title = _get_full_title(data)

    def _empty(size, bold=False):
        p = doc.add_paragraph()
        set_spacing_abs(p, 0, 0)
        r = p.add_run()
        set_run_font(r, FONT_TNR, size, bold=bold, east_asia=FONT_SONG)

    def _centered(text, size, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing_abs(p, 0, 0)
        r = p.add_run(text)
        set_run_font(r, FONT_TNR, size, bold=bold, east_asia=FONT_SONG)

    _empty(PT_4)
    _centered("\u6df1\u5733\u5927\u5b66\u4f20\u64ad\u5b66\u9662", PT_1, bold=True)
    _empty(PT_XIAO1, bold=True)
    _centered("\u7f51\u7edc\u4e0e\u65b0\u5a92\u4f53\u4e13\u4e1a\u672c\u79d1\u6bd5\u4e1a\u8bbe\u8ba1",
              PT_XIAO1, bold=True)
    _centered("\u4e2a\u4eba\u7814\u7a76\u62a5\u544a", PT_XIAO1, bold=True)
    for _ in range(3):
        _empty(PT_3)

    # Title field
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing_abs(p, 0, 0)
    r1 = p.add_run("\u9898\u76ee:")
    set_run_font(r1, FONT_TNR, PT_3, bold=True, east_asia=FONT_SONG)
    r2 = p.add_run(full_title)
    set_run_font(r2, FONT_TNR, PT_3, bold=True, east_asia=FONT_SONG)

    _empty(PT_3, bold=True)

    # Name
    p = doc.add_paragraph()
    set_spacing_abs(p, 0, 0)
    r = p.add_run("\u59d3\u540d:  " + data.get("author", ""))
    set_run_font(r, FONT_TNR, PT_3, bold=True, east_asia=FONT_SONG)

    # Indented fields (8 half-width spaces prefix, matching template)
    for label, value in [
        ("\u5b66\u53f7:  ", "2022080006"),
        ("\u5c0f\u7ec4\u9879\u76ee\uff1a", "\u663e\u00b7\u5f71"),
        ("\u6307\u5bfc\u6559\u5e08: ", "\u674e\u6893\u97f3"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing_abs(p, 0, 0)
        r = p.add_run("        " + label + value)
        set_run_font(r, FONT_TNR, PT_3, bold=True, east_asia=FONT_SONG)

    # Filler
    _empty(PT_3, bold=True)
    for _ in range(4):
        _empty(PT_3)

    # Date
    _centered("\u5e74    \u6708    \u65e5", PT_3)

    # Trailing empties to fill page
    for _ in range(8):
        _empty(PT_5)


# Declaration body text (from official template)
DECLARATION_BODY = (
    "\u672c\u4eba\u90d1\u91cd\u58f0\u660e\uff1a\u6240\u5448\u4ea4\u7684"
    "\u6bd5\u4e1a\u8bbe\u8ba1\u4e2a\u4eba\u7814\u7a76\u62a5\u544a\uff0c"
    "\u9898\u76ee\u300a{title}\u300b"
    " \u662f\u672c\u4eba\u5728\u6307\u5bfc\u6559\u5e08\u7684\u6307\u5bfc"
    "\u4e0b\uff0c\u7ed3\u5408\u9879\u76ee\u5c0f\u7ec4\u8bbe\u8ba1\u5de5"
    "\u4f5c\u5b9e\u8df5\uff0c\u72ec\u7acb\u8fdb\u884c\u7814\u7a76\u5de5"
    "\u4f5c\u6240\u53d6\u5f97\u7684\u6210\u679c\u3002\u672c\u62a5\u544a"
    "\u4e0e\u9879\u76ee\u5c0f\u7ec4\u96c6\u4f53\u5b8c\u6210\u4e4b\u8bbe"
    "\u8ba1\u9879\u76ee\u76f8\u914d\u5408\uff0c\u62a5\u544a\u6240\u5f15"
    "\u8ff0\u4e4b\u76f8\u5173\u8d44\u6599\u53ca\u6587\u732e\uff0c\u5747"
    "\u5728\u62a5\u544a\u4e2d\u4ee5\u660e\u786e\u65b9\u5f0f\u8bf4\u660e"
    "\u3002\u672c\u4eba\u5b8c\u5168\u610f\u8bc6\u5230\u672c\u58f0\u660e"
    "\u7684\u6cd5\u5f8b\u7ed3\u679c\u3002"
)


def build_declaration(doc, data):
    """Generate declaration page (pixel-level replica of official template)."""
    full_title = _get_full_title(data)

    def _empty():
        p = doc.add_paragraph()
        set_spacing_abs(p, 0, 0)
        r = p.add_run()
        set_run_font(r, FONT_TNR, PT_5, east_asia=FONT_SONG)

    _empty()
    _empty()

    # Title lines (18pt bold, centered, spacing before/after = 100 twips)
    for text in [
        "\u6df1\u5733\u5927\u5b66\u672c\u79d1\u6bd5\u4e1a\u8bba\u6587"
        "\u4e2a\u4eba\u7814\u7a76\u62a5\u544a",
        "\u8bda \u4fe1 \u58f0 \u660e",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing_abs(p, 100, 100)
        r = p.add_run(text)
        set_run_font(r, FONT_TNR, PT_XIAO2, bold=True, east_asia=FONT_SONG)

    _empty()

    # Body paragraph (14pt, first-line indent 560 twips)
    body_text = DECLARATION_BODY.format(title=full_title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing_abs(p, 100, 100)
    set_first_line_indent_twips(p, 560)
    r = p.add_run(body_text)
    set_run_font(r, FONT_TNR, PT_4, east_asia=FONT_SONG)

    for _ in range(5):
        _empty()

    # Signature (positioned right via left indent)
    p = doc.add_paragraph()
    set_spacing_abs(p, 0, 0)
    set_left_indent_twips(p, 5000)
    r = p.add_run("\u4f5c\u8005\u7b7e\u540d\uff1a")
    set_run_font(r, FONT_TNR, PT_4, east_asia=FONT_SONG)

    _empty()
    _empty()

    # Date
    p = doc.add_paragraph()
    set_spacing_abs(p, 0, 0)
    set_left_indent_twips(p, 5000)
    r = p.add_run("\u65e5\u671f\uff1a    \u5e74   \u6708   \u65e5")
    set_run_font(r, FONT_TNR, PT_4, east_asia=FONT_SONG)


# Appendix 1 body text (from official template)
APPENDIX1_PARAS = [
    "\u4f20\u64ad\u5b66\u9662\u6bd5\u4e1a\u8bbe\u8ba1\u91c7\u53d6\u5c0f"
    "\u7ec4\u9879\u76ee\u65b9\u5f0f\u8fdb\u884c\uff0c\u6bd5\u4e1a\u8bbe"
    "\u8ba1\u6210\u679c\u4ee5\u9879\u76ee\u6587\u672c\u5f62\u5f0f\u63d0"
    "\u4ea4\uff0c\u6839\u636e\u4e13\u4e1a\u4e0d\u540c\uff0c\u6bd5\u4e1a"
    "\u8bbe\u8ba1\u5927\u81f4\u5206\u4e3a\u4ee5\u4e0b\u51e0\u7c7b\uff1a",
    "1.\u5e7f\u544a\u5b66\u4e13\u4e1a\uff1a\u5546\u4e1a\u54c1\u724c\u8425"
    "\u9500\u7b56\u5212\u63a8\u5e7f\u5168\u6848\u6216\u516c\u76ca\u4f20"
    "\u64ad\u4e0e\u521b\u4e1a\u521b\u610f\u7c7b",
    "2.\u65b0\u95fb\u5b66\u4e13\u4e1a\uff1a\u5168\u5a92\u4f53\u62a5\u9053"
    "\u3001\u4e13\u9898\u7eaa\u5f55\u7247\uff08\u7981\u6b62\u5267\u60c5"
    "\u7247\uff09\u3001\u5927\u578b\u7279\u7a3f\u6216\u901a\u8baf\uff08"
    "\u542b\u4e13\u9898\u62a5\u9053\u548c\u6df1\u5ea6\u62a5\u9053\uff09"
    "\u3001\u7f51\u7edc\u7535\u5b50\u671f\u520a\u6216\u76f8\u5173\u7684"
    "\u65b0\u5a92\u4f53\u5185\u5bb9\u3001\u5a92\u4ecb\u7b56\u5212\u3001"
    "\u793e\u4f1a\u8c03\u7814\u62a5\u544a\u3001\u521b\u529e\u5a92\u4f53"
    "\u7b49",
    "3.\u7f51\u7edc\u4e0e\u65b0\u5a92\u4f53\u4e13\u4e1a\uff1a\u5546\u4e1a"
    "\u4f20\u64ad\u7c7b\u9879\u76ee\u3001\u516c\u5171\u4f20\u64ad\u7c7b"
    "\u9879\u76ee\u3001\u6570\u636e\u4e0e\u53ef\u89c6\u5316\u8bbe\u8ba1"
    "\u3001\u81ea\u5a92\u4f53\u5185\u5bb9\u751f\u4ea7\u4e0e\u5e02\u573a"
    "\u8fd0\u8425\u3001\u65b0\u5a92\u4f53\u827a\u672f\u8bbe\u8ba1\u7b49",
    "\u6bd5\u4e1a\u8bbe\u8ba1\u5c0f\u7ec4\u6210\u5458\u5728\u5b8c\u6210"
    "\u6bd5\u4e1a\u8bbe\u8ba1\u7684\u540c\u65f6\uff0c\u6bcf\u540d\u5c0f"
    "\u7ec4\u6210\u5458\u9700\u72ec\u7acb\u64b0\u5199\u4e00\u7bc7\u4e0d"
    "\u5c11\u4e8e7000\u5b57\u7684\u4e2a\u4eba\u7814\u7a76\u62a5\u544a"
    "\u3002\u4e2a\u4eba\u7814\u7a76\u62a5\u544a\u8981\u6c42\u56f4\u7ed5"
    "\u5c0f\u7ec4\u9879\u76ee\u8fdb\u884c\uff0c\u5e76\u4e3b\u8981\u7ed3"
    "\u5408\u4e2a\u4eba\u5206\u5de5\u4e0e\u5c0f\u7ec4\u9879\u76ee\u6709"
    "\u5173\u7684\u67d0\u4e2a\u5177\u4f53\u95ee\u9898\u8fdb\u884c\u9009"
    "\u9898\uff0c\u5c55\u5f00\u7406\u8bba\u5206\u6790\u548c\u8bba\u8bc1"
    "\uff0c\u529b\u4e89\u6709\u81ea\u5df1\u72ec\u5230\u7684\u53d1\u73b0"
    "\u6216\u8005\u65b0\u9896\u7684\u89c2\u70b9\u3002\u4e0e\u81ea\u5df1"
    "\u5c0f\u7ec4\u9879\u76ee\u65e0\u5173\u7684\u9898\u76ee\uff0c\u4e0d"
    "\u53ef\u4f5c\u4e3a\u4e2a\u4eba\u7814\u7a76\u62a5\u544a\u3002\u4e2a"
    "\u4eba\u7814\u7a76\u62a5\u544a\u5185\u5bb9\u987b\u4e25\u683c\u6309"
    "\u7167\u6df1\u5733\u5927\u5b66\u8bba\u6587\u8981\u6c42\u64b0\u5199"
    "\uff0c\u4e0d\u53ef\u5199\u6210\u5de5\u4f5c\u603b\u7ed3\u6216\u8005"
    "\u64cd\u4f5c\u8fc7\u7a0b\uff0c\u4e0d\u53ef\u91cd\u590d\u9879\u76ee"
    "\u6587\u672c\u5185\u5bb9\u3002",
]


def build_appendix(doc):
    """Generate appendix 1 (explanation) + appendix 2 (member table)."""
    # ---- Appendix 1 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing_abs(p, 0, 0)
    r = p.add_run("\u9644 \u5f551")
    set_run_font(r, FONT_TNR, PT_3, bold=True, east_asia=FONT_SONG)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing_abs(p, 0, 0)
    r = p.add_run("\u6df1\u5733\u5927\u5b66\u4f20\u64ad\u5b66\u9662"
                   "\u4e2a\u4eba\u7814\u7a76\u62a5\u544a\u8bf4\u660e")
    set_run_font(r, FONT_TNR, PT_4, bold=True, east_asia=FONT_SONG)

    for text in APPENDIX1_PARAS:
        p = doc.add_paragraph()
        set_spacing_abs(p, 156, 156)
        set_first_line_indent_twips(p, 420)
        r = p.add_run(text)
        set_run_font(r, FONT_TNR, PT_5, east_asia=FONT_SONG)

    # Signature area (right-aligned)
    for _ in range(2):
        p = doc.add_paragraph()
        set_spacing_abs(p, 0, 0)
        r = p.add_run()
        set_run_font(r, FONT_TNR, PT_5, east_asia=FONT_SONG)

    for sig_text in [
        "\u6df1\u5733\u5927\u5b66\u4f20\u64ad\u5b66\u9662",
        "2022\u5e743\u6708",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing_abs(p, 0, 0)
        r = p.add_run(sig_text)
        set_run_font(r, FONT_TNR, PT_5, east_asia=FONT_SONG)

    # ---- Appendix 2 (new page) ----
    add_page_break(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing_abs(p, 0, 0)
    r = p.add_run("\u9644 \u5f552")
    set_run_font(r, FONT_TNR, PT_3, bold=True, east_asia=FONT_SONG)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing_abs(p, 0, 0)
    r = p.add_run("\u300a\u663e\u00b7\u5f71\u300b"
                   "\u9879\u76ee\u5c0f\u7ec4\u6210\u5458\u5206\u5de5\u8bf4\u660e")
    set_run_font(r, FONT_TNR, PT_4, bold=True, east_asia=FONT_SONG)

    # 8 rows (1 header + 7 data), 3 columns
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(8.5)

    headers = ["\u5e8f\u53f7", "\u5b66\u53f7/\u59d3\u540d",
               "\u5206\u5de5\u3001\u5de5\u4f5c\u5185\u5bb9"]
    for j, text in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        tc_pr = cell._tc.get_or_add_tcPr()
        va = OxmlElement("w:vAlign")
        va.set(qn("w:val"), "center")
        tc_pr.append(va)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing_abs(p, 0, 0)
        r = p.add_run(text)
        set_run_font(r, FONT_TNR, PT_5, east_asia=FONT_SONG)

    for i in range(7):
        cell = table.rows[i + 1].cells[0]
        cell.text = ""
        tc_pr = cell._tc.get_or_add_tcPr()
        va = OxmlElement("w:vAlign")
        va.set(qn("w:val"), "center")
        tc_pr.append(va)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing_abs(p, 0, 0)
        r = p.add_run(str(i + 1))
        set_run_font(r, FONT_TNR, PT_5, east_asia=FONT_SONG)
        for j in range(1, 3):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            tc_pr = cell._tc.get_or_add_tcPr()
            va = OxmlElement("w:vAlign")
            va.set(qn("w:val"), "center")
            tc_pr.append(va)
            p = cell.paragraphs[0]
            set_spacing_abs(p, 0, 0)


# ---------------------------------------------------------------------------
# Figure counter for mermaid images
# ---------------------------------------------------------------------------
_fig_counter = 0


def next_fig_num():
    global _fig_counter
    _fig_counter += 1
    return _fig_counter


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main():
    data = parse_md(MD_FILE)
    doc = Document()

    # Section 1: Cover (wider margins matching official template)
    _set_section_layout(doc.sections[0])
    build_cover(doc, data)

    # Section 2: Declaration (same margins)
    sec1 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_section_layout(sec1)
    build_declaration(doc, data)

    # Section 3: Main content (2.5cm all sides per SZU spec)
    sec2 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_section_layout(sec2, top=2.5, bottom=2.5, left=2.5, right=2.5)
    build_toc(doc)
    add_page_break(doc)
    build_title_page(doc, data)
    build_abstract(doc, data["abstract"], data["keywords"])

    pending_label = None
    for item in data["body_lines"]:
        kind = item[0]
        if kind == "h1":
            build_heading(doc, 1, item[1])
        elif kind == "h2":
            build_heading(doc, 2, item[1])
        elif kind == "h3":
            build_heading(doc, 3, item[1])
        elif kind == "para":
            build_body_paragraph(doc, item[1])
        elif kind == "note":
            build_note_paragraph(doc, item[1])
        elif kind == "img_label":
            pending_label = item[1]
        elif kind == "mermaid":
            fig_num = next_fig_num()
            img_path = DIR / IMAGE_MAP.get(fig_num, "")
            if img_path.exists():
                build_image(doc, img_path)
                if pending_label:
                    build_img_label(doc, pending_label)
                    pending_label = None
            else:
                if pending_label:
                    build_img_label(doc, pending_label
                                    + "\uff08Mermaid\u56fe\u7247\u7f3a\u5931\uff09")
                    pending_label = None
        elif kind == "fig4_html":
            pending_label = None
            build_fig4(doc)
        elif kind == "table":
            build_table(doc, item[1])
            if pending_label:
                build_img_label(doc, pending_label)
                pending_label = None

    build_references(doc, data["references"])
    build_thanks(doc, data["thanks"])
    build_en_abstract(doc, data["en_abstract"], data["en_keywords"])

    # Section 4: Appendix (wider margins)
    sec3 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_section_layout(sec3)
    build_appendix(doc)

    doc.save(str(OUT_FILE))
    print(f"Generated: {OUT_FILE}")
    print(f"  Sections: {len(doc.sections)}")


if __name__ == "__main__":
    main()
